"""
文档处理服务 - 提取各类文件内容
"""
import io
import re
from typing import Tuple


class DocumentProcessor:
    """文档内容提取器"""

    SUPPORTED_TYPES = {
        'pdf': ['application/pdf'],
        'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        'txt': ['text/plain'],
        'md': ['text/markdown', 'text/plain'],
        'xlsx': ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']
    }

    @classmethod
    def get_file_type(cls, filename: str) -> str:
        """根据文件名获取文件类型"""
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        return ext if ext in cls.SUPPORTED_TYPES else 'unknown'

    @classmethod
    def extract_content(cls, file_content: bytes, file_type: str, filename: str) -> Tuple[str, str]:
        """提取文件内容（从字节）"""
        type_names = {
            'pdf': 'PDF', 'docx': 'DOCX', 'txt': 'TXT', 'md': 'MD', 'xlsx': 'XLSX'
        }

        if file_type == 'unknown':
            raise ValueError(f"不支持的文件类型: {file_type}")

        if file_type == 'pdf':
            content = cls._extract_pdf(file_content)
        else:
            extractors = {
                'docx': cls._extract_docx,
                'txt': cls._extract_txt,
                'md': cls._extract_txt,
                'xlsx': cls._extract_xlsx
            }
            extractor = extractors.get(file_type)
            if not extractor:
                raise ValueError(f"不支持的文件类型: {file_type}")
            content = extractor(file_content)

        return content, type_names.get(file_type, file_type.upper())

    @classmethod
    def extract_from_file(cls, file_path: str, file_type: str, filename: str) -> Tuple[str, str]:
        """提取文件内容（从文件路径，流式读取避免内存爆）"""
        type_names = {
            'pdf': 'PDF', 'docx': 'DOCX', 'txt': 'TXT', 'md': 'MD', 'xlsx': 'XLSX'
        }

        if file_type == 'unknown':
            raise ValueError(f"不支持的文件类型: {file_type}")

        if file_type == 'pdf':
            content = cls._extract_pdf_from_path(file_path)
        elif file_type in ('txt', 'md'):
            content = cls._extract_txt_from_path(file_path)
        elif file_type == 'docx':
            content = cls._extract_docx_from_path(file_path)
        elif file_type == 'xlsx':
            content = cls._extract_xlsx_from_path(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

        return content, type_names.get(file_type, file_type.upper())

    @staticmethod
    def _fix_pdf_math_chars(text: str) -> str:
        """修复PDF数学字体提取乱码 - 将错误映射的CJK字符还原为数学符号

        人教版数学教材PDF使用自定义Type1字体，将Latin字母映射到了CJK码点。
        该映射基于对多本人教版教材的分析，覆盖 E-BX9 (斜体)、E-BZ9 (正体) 等字体。
        """
        # === 斜体数学变量字体 (E-BX9) 的CJK→Latin映射 ===
        cjk_to_latin = {
            # 小写字母
            '犪': 'a', '犫': 'b', '犮': 'c', '犱': 'd', '犲': 'e',
            '犳': 'f', '犺': 'h', '犻': 'i', '犽': 'k', '犾': 'l',
            '犿': 'm', '狀': 'n', '狆': 'p', '狇': 'q', '狉': 'r',
            '狊': 's', '狋': 't', '狓': 'x', '狔': 'y', '狌': 'u',
            '犼': 'v', '狅': 'o', '犼': 'w', '犰': 'j', '犷': 'g',
            '犿': 'm',
            # 大写字母
            '犃': 'A', '犅': 'B', '犆': 'C', '犇': 'D', '犈': 'E',
            '犉': 'F', '犌': 'G', '犎': 'H', '犐': 'I', '犑': 'J',
            '犓': 'K', '犔': 'L', '犕': 'M', '犖': 'N', '犗': 'O',
            '犘': 'P', '犙': 'Q', '犚': 'R', '犛': 'S', '犜': 'T',
            '犝': 'U', '犞': 'V', '犠': 'W', '犡': 'X', '犢': 'Y',
            '犣': 'Z',
            # 方程组大括号
            '烄': '{', '烆': '}', '烌': '}',
            # 条件概率竖线
            '狘': '|',
            # 数学符号
            '槿': '√', '槡': '√',
        }
        for wrong, correct in cjk_to_latin.items():
            text = text.replace(wrong, correct)

        # === 全角字符→半角 ===
        text = DocumentProcessor._fullwidth_to_halfwidth(text)

        # === Private Use Area 字符映射 ===
        # 数学教材PUA字符
        pua_map = {
            '\ue010': '.',   # 分隔符 6.1
            '\ue011': '-',   # 连字符 6.1-1
            '\ue012': '*',   # 星号 N*
            '\ue020': '⊂',   # 子集
            '\ue03c': 'Σ',   # 求和
            '\ue05b': '∉',   # 不属于
            '\ue07e': '∅',   # 空集
        }
        # 英语教材音标字体PUA字符 (U+F0xx → IPA音标)
        # 人教版英语教材PDF使用自定义音标字体，IPA符号编码为 ASCII码点+0xF000
        pua_map.update({
            '\uf022': 'ˈ',   # 主重音 "
            '\uf025': 'ˌ',   # 次重音 %
            '\uf02f': '/',   # 音标边界 /
            '\uf033': 'ɜ',   # open-mid central 3
            '\uf03a': 'ː',   # 长音标记 :
            '\uf040': 'ə',   # schwa @
            '\uf041': 'ɑ',   # open back A
            '\uf044': 'ð',   # eth D
            '\uf049': 'ɪ',   # near-close front I
            '\uf04e': 'ŋ',   # eng N
            '\uf04f': 'ɔ',   # open-mid back O
            '\uf051': 'ɒ',   # open back rounded Q
            '\uf053': 'ʃ',   # esh S
            '\uf054': 'θ',   # theta T
            '\uf055': 'ʊ',   # near-close back U
            '\uf056': 'ʌ',   # open-mid back V
            '\uf05a': 'ʒ',   # ezh Z
            '\uf07b': 'æ',   # ash {
        })
        # 英语教材音标字体中的小写字母 a-z（直接还原为自身）
        for _c in range(ord('a'), ord('z') + 1):
            pua_key = chr(0xF000 + _c)
            if pua_key not in pua_map:
                pua_map[pua_key] = chr(_c)
        # 音标中的大写字母和数字也直接还原
        for _c in list(range(ord('A'), ord('Z') + 1)) + list(range(ord('0'), ord('9') + 1)):
            pua_key = chr(0xF000 + _c)
            if pua_key not in pua_map:
                pua_map[pua_key] = chr(_c)
        # 音标中的标点符号也还原
        for _c in ' ()-.,;\'!':
            pua_key = chr(0xF000 + ord(_c))
            if pua_key not in pua_map:
                pua_map[pua_key] = _c
        for wrong, correct in pua_map.items():
            text = text.replace(wrong, correct)

        # 清除剩余的PUA字符（无法识别的装饰符号）
        text = re.sub(r'[\ue000-\uf8ff]', '', text)

        return text

    @staticmethod
    def _fullwidth_to_halfwidth(text: str) -> str:
        """将全角字母/数字/符号转为半角，保留中文标点"""
        result = []
        for ch in text:
            code = ord(ch)
            # 全角数字 ０-９ → 0-9
            if 0xFF10 <= code <= 0xFF19:
                result.append(chr(code - 0xFEE0))
            # 全角大写 A-Z → A-Z
            elif 0xFF21 <= code <= 0xFF3A:
                result.append(chr(code - 0xFEE0))
            # 全角小写 a-z → a-z
            elif 0xFF41 <= code <= 0xFF5A:
                result.append(chr(code - 0xFEE0))
            # 全角符号
            elif ch == '＝':
                result.append('=')
            elif ch == '＋':
                result.append('+')
            elif ch == '－':
                result.append('-')
            elif ch == '×':
                result.append('×')  # 保留乘号
            elif ch == '÷':
                result.append('÷')  # 保留除号
            elif ch == '．':
                result.append('.')
            elif ch == '（':
                result.append('(')
            elif ch == '）':
                result.append(')')
            elif ch == '，':
                result.append(',')
            elif ch == '：':
                result.append(':')
            elif ch == '；':
                result.append(';')
            elif ch == '！':
                result.append('!')
            elif ch == '？':
                result.append('?')
            elif ch == '＜':
                result.append('<')
            elif ch == '＞':
                result.append('>')
            elif ch == '≤':
                result.append('≤')
            elif ch == '≥':
                result.append('≥')
            elif ch == '≠':
                result.append('≠')
            else:
                result.append(ch)
        return ''.join(result)

    @staticmethod
    def _deduplicate_overlapping_chars(text: str) -> str:
        """去除PDF封面/水印层导致的重复字符

        例如 '绿绿色色印印刷刷' → '绿色印刷', '普\n通 普 通 高 中' → '普通高中'
        """
        lines = text.split('\n')
        result = []
        for line in lines:
            # 检测连续重复字符对: "绿绿" "色色" "印印"
            deduped = re.sub(r'(.)\1', r'\1', line)
            # 检测空格分隔的重复: "普 通 高 中" 前后行有 "普通高中"
            # 这个在后续行合并时处理
            result.append(deduped)
        return '\n'.join(result)

    @staticmethod
    def _is_junk_page(text: str) -> bool:
        """判断是否为PDF的封面、版权页等非内容页面"""
        if not text or not text.strip():
            return True
        # 封面/版权页特征
        junk_markers = [
            '定价', '版权所有', '未经许可', 'ISBN',
            '责任编辑', '美术编辑', '封面设计', '版面设计',
            '印刷', '出版', '邮编',
        ]
        text_stripped = text.strip()
        # 很短的页面（只有几行）
        lines = [l.strip() for l in text_stripped.split('\n') if l.strip()]
        if len(lines) <= 3:
            return True

        # 包含多个版权特征的页面
        marker_count = sum(1 for m in junk_markers if m in text_stripped)
        if marker_count >= 2:
            return True

        # 高密度重复字符（水印层叠加）
        if len(text_stripped) > 20:
            deduped = re.sub(r'(.)\1', r'\1', text_stripped)
            if len(deduped) < len(text_stripped) * 0.6:
                return True

        return False

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        """提取PDF内容，对数学教材自动使用VL模型增强"""
        import pdfplumber

        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if not text:
                    continue
                # 跳过封面/版权页
                if DocumentProcessor._is_junk_page(text):
                    continue
                # 修复数学字体乱码
                text = DocumentProcessor._fix_pdf_math_chars(text)
                # 去除重复字符
                text = DocumentProcessor._deduplicate_overlapping_chars(text)
                # 清理多余空白
                text = re.sub(r' {3,}', ' ', text)
                text_parts.append(text)

        extracted = '\n\n'.join(text_parts)

        # 如果pdfplumber几乎没有提取到内容（扫描版PDF），用VL模型识别多页
        if len(extracted.strip()) < 100:
            try:
                from services.vl_pdf_extractor import extract_scanned_pdf_with_vl
                vl_text = extract_scanned_pdf_with_vl(content)
                if vl_text:
                    return f"[VL模型识别]\n{vl_text}"
            except Exception as e:
                import logging
                logging.getLogger(__name__).warning(f"扫描版PDF VL提取失败: {e}")
            return extracted

        # 检测是否有乱码残留，如果有则用VL模型增强
        garbled_count = sum(1 for ch in extracted if 0x7280 <= ord(ch) <= 0x72E0)
        if garbled_count > 5:
            try:
                from services.vl_pdf_extractor import extract_pdf_with_vl
                vl_text = extract_pdf_with_vl(content)
                if vl_text:
                    extracted = f"{vl_text}\n\n{extracted}"
            except Exception as e:
                import logging
                logging.getLogger(__name__).warning(f"VL增强失败，使用原始提取: {e}")

        return extracted

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        """提取DOCX内容"""
        from docx import Document

        doc = Document(io.BytesIO(content))
        return '\n'.join([para.text for para in doc.paragraphs if para.text])

    @staticmethod
    def _extract_txt(content: bytes) -> str:
        """提取TXT/MD内容"""
        candidates = ['utf-8-sig', 'utf-8', 'gb18030', 'gbk', 'gb2312']

        def looks_garbled(text: str) -> bool:
            if not text:
                return False
            bad_chars = text.count('\ufffd')
            mojibake_markers = text.count('锟') + text.count('鈥') + text.count('Ã')
            return bad_chars > 0 or mojibake_markers > 0

        for encoding in candidates:
            try:
                decoded = content.decode(encoding)
                if not looks_garbled(decoded):
                    return decoded
            except UnicodeDecodeError:
                continue

        decoded = content.decode('utf-8', errors='replace')
        if not looks_garbled(decoded):
            return decoded

        return content.decode('gb18030', errors='replace')

    @staticmethod
    def _extract_xlsx(content: bytes) -> str:
        """提取XLSX内容"""
        import pandas as pd

        df_dict = pd.read_excel(io.BytesIO(content), sheet_name=None)
        text_parts = []
        for sheet_name, sheet_df in df_dict.items():
            text_parts.append(f"=== 工作表: {sheet_name} ===")
            text_parts.append(sheet_df.to_string())
        return '\n\n'.join(text_parts)

    # ---- 从文件路径提取（流式，省内存） ----

    @staticmethod
    def _extract_pdf_from_path(file_path: str) -> str:
        """从文件路径提取PDF内容"""
        import pdfplumber

        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if not text:
                    continue
                if DocumentProcessor._is_junk_page(text):
                    continue
                text = DocumentProcessor._fix_pdf_math_chars(text)
                text = DocumentProcessor._deduplicate_overlapping_chars(text)
                text = re.sub(r' {3,}', ' ', text)
                text_parts.append(text)

        extracted = '\n\n'.join(text_parts)

        # 扫描版PDF
        if len(extracted.strip()) < 100:
            try:
                from services.vl_pdf_extractor import extract_scanned_pdf_with_vl_from_path
                vl_text = extract_scanned_pdf_with_vl_from_path(file_path)
                if vl_text:
                    return f"[VL模型识别]\n{vl_text}"
            except Exception as e:
                import logging
                logging.getLogger(__name__).warning(f"扫描版PDF VL提取失败: {e}")
            return extracted

        # 乱码增强
        garbled_count = sum(1 for ch in extracted if 0x7280 <= ord(ch) <= 0x72E0)
        if garbled_count > 5:
            try:
                from services.vl_pdf_extractor import extract_pdf_with_vl_from_path
                vl_text = extract_pdf_with_vl_from_path(file_path)
                if vl_text:
                    extracted = f"{vl_text}\n\n{extracted}"
            except Exception as e:
                import logging
                logging.getLogger(__name__).warning(f"VL增强失败，使用原始提取: {e}")

        return extracted

    @staticmethod
    def _extract_txt_from_path(file_path: str) -> str:
        """从文件路径提取TXT/MD内容"""
        with open(file_path, "rb") as f:
            content = f.read()
        return DocumentProcessor._extract_txt(content)

    @staticmethod
    def _extract_docx_from_path(file_path: str) -> str:
        """从文件路径提取DOCX内容"""
        from docx import Document
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs if para.text])

    @staticmethod
    def _extract_xlsx_from_path(file_path: str) -> str:
        """从文件路径提取XLSX内容"""
        import pandas as pd
        df_dict = pd.read_excel(file_path, sheet_name=None)
        text_parts = []
        for sheet_name, sheet_df in df_dict.items():
            text_parts.append(f"=== 工作表: {sheet_name} ===")
            text_parts.append(sheet_df.to_string())
        return '\n\n'.join(text_parts)
